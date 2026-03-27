from docx import Document
from docx.shared import Pt


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_heading(doc, "Low-Level Architecture - Bot Soberania (SoftwareOne)", 0)
    add_para(doc, "Projeto: C:/Users/NATH/Downloads/Bot_Soberania")
    add_para(
        doc,
        "Repositorio GitHub: https://github.com/Akflamingo2610/BotSoberaniaSoftwareOne",
    )
    doc.add_paragraph()

    add_heading(doc, "1. Visao Geral de Componentes", 1)
    add_para(
        doc,
        "A solucao e composta por um app Flutter (frontend), API de negocio no Xano, "
        "e um servidor RAG em Node.js para respostas contextuais e explicacoes em streaming.",
    )
    add_bullets(
        doc,
        [
            "soberania_app/: aplicativo Flutter (web/mobile).",
            "Xano API: autenticacao, questoes, progresso e salvamento de respostas.",
            "rag_server/: backend RAG para chat e explicacoes automáticas.",
            "rag_server/docs/: PDFs indexados para busca semantica aproximada.",
        ],
    )

    add_heading(doc, "2. Topologia de Runtime", 1)
    add_para(
        doc,
        "Fluxo principal de rede:",
    )
    add_para(
        doc,
        "[Flutter App] -> [Xano API] (assessment)\n"
        "[Flutter App] -> [RAG Server] (chat e explicacoes)\n"
        "[RAG Server] -> [Groq API ou Ollama] (geracao de resposta)",
    )

    add_heading(doc, "3. Flutter - Estrutura Low-Level", 1)
    add_bullets(
        doc,
        [
            "main.dart: bootstrap, tema, locale e MaterialApp.",
            "l10n/: AppLocalizations e LocaleScope (PT/EN/ES).",
            "config.dart: xanoBaseUrl e ragBaseUrl.",
            "api/xano_api.dart: client HTTP + cache (questoes/progresso).",
            "api/rag_api.dart: endpoints de chat, stream e health.",
            "storage/app_storage.dart: persistencia local com SharedPreferences.",
            "models/models.dart: Question, SavedAnswer e normalizacao de score.",
            "screens/: fluxo de telas (welcome, login, intro, fases, questoes, resultados).",
            "widgets/chat_panel.dart: painel de chat com streaming e estado online/offline.",
        ],
    )

    add_heading(doc, "4. Fluxo Funcional - Assessment", 1)
    add_bullets(
        doc,
        [
            "Login -> recebe token do Xano.",
            "Resume assessment -> obtencao/reativacao do assessment_id.",
            "Listagem de questoes por fase ou pilar.",
            "Carga de progresso para pre-preencher respostas.",
            "Resposta local pendente e salvamento em lote no endpoint /assessment/save.",
            "Calculo de resultados por pilar e dominio na tela de Results.",
        ],
    )

    add_heading(doc, "5. Fluxo Funcional - Chat/RAG", 1)
    add_bullets(
        doc,
        [
            "Health check inicial do RAG.",
            "Perguntas do usuario via /ask ou /ask/stream.",
            "Explicacao automatica de pergunta via /ask/explain-question/stream.",
            "Explicacao em lote via /ask/explain-batch.",
            "Fallback de conectividade: status online e atualizado por sucesso real de resposta.",
        ],
    )

    add_heading(doc, "6. RAG Server - Arquitetura Interna", 1)
    add_bullets(
        doc,
        [
            "Leitura de PDFs com pdf-parse.",
            "Chunking de texto para indexacao.",
            "Indice local com MiniSearch (titulo + texto).",
            "Heuristica para priorizar fontes AWS vs leis.",
            "Prompt builder com idioma (pt/en/es) e regras de relevancia.",
            "LLM primario: Groq; fallback: Ollama (quando habilitado).",
            "Saida em JSON ou NDJSON para streaming.",
        ],
    )

    add_heading(doc, "7. Endpoints Principais", 1)
    add_para(doc, "Xano (consumido pelo app):", bold=True)
    add_bullets(
        doc,
        [
            "POST /login",
            "POST /signup_company",
            "POST /assessment/resume",
            "GET /questions?phase={phase}",
            "GET /progress/assessment?assessment_id={id}",
            "POST /assessment/save",
            "POST /forgot_password",
            "POST /reset_password",
        ],
    )
    add_para(doc, "RAG Server:", bold=True)
    add_bullets(
        doc,
        [
            "GET /health",
            "POST /ask",
            "POST /ask/stream",
            "POST /ask/explain-question/stream",
            "POST /ask/explain-batch",
        ],
    )

    add_heading(doc, "8. Persistencia e Estado", 1)
    add_bullets(
        doc,
        [
            "authToken, assessmentId, userName, userEmail.",
            "locale (idioma atual).",
            "introSeen (controle de onboarding).",
            "lastQuestionIndex por fase (retomada de navegacao).",
            "cache em memoria para reduzir chamadas ao Xano.",
        ],
    )

    add_heading(doc, "9. Deploy e Operacao", 1)
    add_bullets(
        doc,
        [
            "Flutter Web: build + Firebase Hosting.",
            "RAG Server: Railway (prod) ou local em localhost:4000.",
            "Variaveis criticas: xanoBaseUrl, ragBaseUrl, GROQ_API_KEY, USE_OLLAMA.",
            "Cold start no Railway pode impactar a primeira resposta.",
        ],
    )

    add_heading(doc, "10. Pontos de Atencao Tecnicos", 1)
    add_bullets(
        doc,
        [
            "Consistencia de idioma entre UI, perguntas e respostas do chat.",
            "Tratamento de falhas temporarias do /health sem falso offline.",
            "Controle de navegacao no modo 1-card vs bloco de 9 perguntas.",
            "Manter documentos RAG atualizados para respostas relevantes.",
        ],
    )

    output_path = "LOW_LEVEL_ARCHITECTURE_BotSoberania.docx"
    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
