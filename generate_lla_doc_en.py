from docx import Document
from docx.shared import Pt


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    return p


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_heading(doc, "Low-Level Architecture - Bot Soberania (SoftwareOne)", 0)
    add_para(doc, "Project: C:/Users/NATH/Downloads/Bot_Soberania")
    add_para(
        doc,
        "GitHub Repository: https://github.com/Akflamingo2610/BotSoberaniaSoftwareOne",
    )
    doc.add_paragraph()

    add_heading(doc, "1. Component Overview", 1)
    add_para(
        doc,
        "The solution is composed of a Flutter app (frontend), a business API on Xano, "
        "and a Node.js RAG server for contextual responses and streaming explanations.",
    )
    add_bullets(
        doc,
        [
            "soberania_app/: Flutter app (web/mobile).",
            "Xano API: authentication, questions, progress and answer persistence.",
            "rag_server/: RAG backend for chat and automatic explanations.",
            "rag_server/docs/: indexed PDFs for search/retrieval.",
        ],
    )

    add_heading(doc, "2. Runtime Topology", 1)
    add_para(doc, "Main network flow:")
    add_para(
        doc,
        "[Flutter App] -> [Xano API] (assessment)\n"
        "[Flutter App] -> [RAG Server] (chat and explanations)\n"
        "[RAG Server] -> [Groq API or Ollama] (response generation)",
    )

    add_heading(doc, "3. Flutter - Low-Level Structure", 1)
    add_bullets(
        doc,
        [
            "main.dart: bootstrap, theme, locale and MaterialApp.",
            "l10n/: AppLocalizations and LocaleScope (PT/EN/ES).",
            "config.dart: xanoBaseUrl and ragBaseUrl.",
            "api/xano_api.dart: HTTP client + in-memory cache (questions/progress).",
            "api/rag_api.dart: chat, stream and health endpoints.",
            "storage/app_storage.dart: local persistence with SharedPreferences.",
            "models/models.dart: Question, SavedAnswer and score normalization.",
            "screens/: screen flow (welcome, login, intro, phases, questions, results).",
            "widgets/chat_panel.dart: chat panel with streaming and online/offline state.",
        ],
    )

    add_heading(doc, "4. Functional Flow - Assessment", 1)
    add_bullets(
        doc,
        [
            "Login -> receives token from Xano.",
            "Resume assessment -> obtains/reactivates assessment_id.",
            "Loads questions by phase or pillar.",
            "Loads progress to prefill saved answers.",
            "Keeps local pending answers and persists in bulk through /assessment/save.",
            "Computes results by pillar and domain on Results screen.",
        ],
    )

    add_heading(doc, "5. Functional Flow - Chat/RAG", 1)
    add_bullets(
        doc,
        [
            "Initial RAG health check.",
            "User questions via /ask or /ask/stream.",
            "Question auto-explanation via /ask/explain-question/stream.",
            "Batch explanation via /ask/explain-batch.",
            "Connectivity fallback: online status also updated by actual successful responses.",
        ],
    )

    add_heading(doc, "6. RAG Server - Internal Architecture", 1)
    add_bullets(
        doc,
        [
            "Reads PDFs with pdf-parse.",
            "Text chunking for indexing.",
            "Local index with MiniSearch (title + text).",
            "Heuristics to prioritize AWS docs vs legal docs.",
            "Prompt builder with language (pt/en/es) and relevance rules.",
            "Primary LLM: Groq; fallback: Ollama (if enabled).",
            "Output format: JSON or NDJSON streaming.",
        ],
    )

    add_heading(doc, "7. Main Endpoints", 1)
    add_para(doc, "Xano (consumed by app):", bold=True)
    add_bullets(
        doc,
        [
            "POST /login",
            "POST /signup_company",
            "POST /assessment/resume",
            "GET /questions?phase={phase}",
            "GET /progress/assessment?assessment_id={id}",
            "POST /assessment/save",
            "POST /forgot_password",
            "POST /reset_password",
        ],
    )
    add_para(doc, "RAG Server:", bold=True)
    add_bullets(
        doc,
        [
            "GET /health",
            "POST /ask",
            "POST /ask/stream",
            "POST /ask/explain-question/stream",
            "POST /ask/explain-batch",
        ],
    )

    add_heading(doc, "8. Persistence and State", 1)
    add_bullets(
        doc,
        [
            "authToken, assessmentId, userName, userEmail.",
            "locale (current language).",
            "introSeen (onboarding control).",
            "lastQuestionIndex per phase (resume support).",
            "in-memory cache to reduce Xano requests.",
        ],
    )

    add_heading(doc, "9. Deployment and Operations", 1)
    add_bullets(
        doc,
        [
            "Flutter Web: build + Firebase Hosting.",
            "RAG Server: Railway (prod) or localhost:4000 (local).",
            "Critical variables: xanoBaseUrl, ragBaseUrl, GROQ_API_KEY, USE_OLLAMA.",
            "Railway cold starts may affect first response latency.",
        ],
    )

    add_heading(doc, "10. Technical Attention Points", 1)
    add_bullets(
        doc,
        [
            "Language consistency across UI, questions and chat responses.",
            "Avoid false offline states on temporary /health failures.",
            "Correct navigation behavior in 1-card mode vs 9-question block mode.",
            "Keep RAG documents updated for relevant answers.",
        ],
    )

    output_path = "LOW_LEVEL_ARCHITECTURE_BotSoberania_EN.docx"
    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
